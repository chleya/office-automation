"""
Word Document Processor

This module provides functionality for creating, editing, and manipulating
Word documents (.docx format) with WPS compatibility.
"""

import os
from typing import Optional, List, Dict, Any, Union
from pathlib import Path

from ..utils.error_handler import (
    DocumentCreationError,
    DocumentReadError,
    DocumentSaveError,
    ValidationError,
    create_error_context,
    validate_file_path,
    validate_format,
)


class WordProcessor:
    """Processor for Word document operations."""
    
    # Supported formats
    SUPPORTED_FORMATS = ["docx", "doc", "pdf", "rtf", "txt"]
    
    def __init__(self, template_dir: Optional[str] = None):
        """
        Initialize Word Processor.
        
        Args:
            template_dir: Directory containing Word templates
        """
        self.template_dir = template_dir
        self._document = None
        
        # Try to import python-docx
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            self.Document = Document
            self.Inches = Inches
            self.Pt = Pt
            self.RGBColor = RGBColor
            self.WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
            self.WD_STYLE_TYPE = WD_STYLE_TYPE
            
            self._docx_available = True
        except ImportError:
            self._docx_available = False
            raise ImportError(
                "python-docx is not installed. Please install it with: "
                "pip install python-docx"
            )
    
    def create_document(
        self,
        template: Optional[Union[str, Path]] = None
    ) -> 'Document':
        """
        Create a new Word document.
        
        Args:
            template: Path to template file or template name
            
        Returns:
            Document object
            
        Raises:
            DocumentCreationError: If document creation fails
        """
        context = create_error_context("create_document", template=template)
        
        try:
            if template:
                # Load from template
                template_path = self._resolve_template_path(template)
                if not os.path.exists(template_path):
                    raise FileNotFoundError(f"Template not found: {template_path}")
                
                self._document = self.Document(template_path)
            else:
                # Create empty document
                self._document = self.Document()
            
            return self._document
            
        except Exception as e:
            raise DocumentCreationError(
                f"Failed to create document: {e}",
                context
            )
    
    def load_document(self, file_path: Union[str, Path]) -> 'Document':
        """
        Load an existing Word document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object
            
        Raises:
            DocumentReadError: If document loading fails
        """
        context = create_error_context("load_document", file_path=file_path)
        validate_file_path(str(file_path), "load_document")
        
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            self._document = self.Document(str(file_path))
            return self._document
            
        except Exception as e:
            raise DocumentReadError(
                f"Failed to load document: {e}",
                context
            )
    
    def add_heading(
        self,
        text: str,
        level: int = 1,
        style: Optional[str] = None
    ) -> None:
        """
        Add a heading to the document.
        
        Args:
            text: Heading text
            level: Heading level (1-9)
            style: Custom style name
            
        Raises:
            DocumentCreationError: If document is not initialized
        """
        if not self._document:
            raise DocumentCreationError(
                "Document not initialized. Call create_document() or load_document() first."
            )
        
        try:
            heading = self._document.add_heading(text, level)
            if style:
                heading.style = style
        except Exception as e:
            raise DocumentCreationError(f"Failed to add heading: {e}")
    
    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        alignment: Optional[str] = None
    ) -> None:
        """
        Add a paragraph to the document.
        
        Args:
            text: Paragraph text
            style: Paragraph style
            alignment: Text alignment (left, center, right, justify)
        """
        if not self._document:
            raise DocumentCreationError(
                "Document not initialized. Call create_document() or load_document() first."
            )
        
        try:
            paragraph = self._document.add_paragraph(text)
            
            if style:
                paragraph.style = style
            
            if alignment:
                align_map = {
                    "left": self.WD_ALIGN_PARAGRAPH.LEFT,
                    "center": self.WD_ALIGN_PARAGRAPH.CENTER,
                    "right": self.WD_ALIGN_PARAGRAPH.RIGHT,
                    "justify": self.WD_ALIGN_PARAGRAPH.JUSTIFY,
                }
                if alignment in align_map:
                    paragraph.alignment = align_map[alignment]
                    
        except Exception as e:
            raise DocumentCreationError(f"Failed to add paragraph: {e}")
    
    def add_table(
        self,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        style: Optional[str] = "Table Grid"
    ) -> None:
        """
        Add a table to the document.
        
        Args:
            data: Table data as list of rows
            headers: Column headers
            style: Table style
        """
        if not self._document:
            raise DocumentCreationError(
                "Document not initialized. Call create_document() or load_document() first."
            )
        
        try:
            # Calculate table dimensions
            num_rows = len(data)
            num_cols = len(data[0]) if data else 0
            
            if headers:
                num_rows += 1
            
            if num_rows == 0 or num_cols == 0:
                return
            
            # Create table
            table = self._document.add_table(rows=num_rows, cols=num_cols)
            
            # Apply style
            if style:
                table.style = style
            
            # Add headers if provided
            if headers:
                header_row = table.rows[0]
                for i, header in enumerate(headers):
                    cell = header_row.cells[i]
                    cell.text = str(header)
                    # Make header bold
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            
            # Add data
            start_row = 1 if headers else 0
            for i, row in enumerate(data):
                table_row = table.rows[start_row + i]
                for j, cell_value in enumerate(row):
                    cell = table_row.cells[j]
                    cell.text = str(cell_value)
                    
        except Exception as e:
            raise DocumentCreationError(f"Failed to add table: {e}")
    
    def add_image(
        self,
        image_path: Union[str, Path],
        width: Optional[float] = None,
        height: Optional[float] = None,
        caption: Optional[str] = None
    ) -> None:
        """
        Add an image to the document.
        
        Args:
            image_path: Path to image file
            width: Image width in inches
            height: Image height in inches
            caption: Image caption
        """
        if not self._document:
            raise DocumentCreationError(
                "Document not initialized. Call create_document() or load_document() first."
            )
        
        try:
            image_path = Path(image_path)
            if not image_path.exists():
                raise FileNotFoundError(f"Image not found: {image_path}")
            
            # Add image
            if width and height:
                self._document.add_picture(
                    str(image_path),
                    width=self.Inches(width),
                    height=self.Inches(height)
                )
            elif width:
                self._document.add_picture(str(image_path), width=self.Inches(width))
            elif height:
                self._document.add_picture(str(image_path), height=self.Inches(height))
            else:
                self._document.add_picture(str(image_path))
            
            # Add caption if provided
            if caption:
                self.add_paragraph(caption, style="Caption")
                
        except Exception as e:
            raise DocumentCreationError(f"Failed to add image: {e}")
    
    def save(
        self,
        file_path: Union[str, Path],
        format: str = "docx"
    ) -> str:
        """
        Save the document to a file.
        
        Args:
            file_path: Path where to save the document
            format: Output format (docx, pdf, etc.)
            
        Returns:
            Path to saved file
            
        Raises:
            DocumentSaveError: If saving fails
        """
        context = create_error_context(
            "save_document",
            file_path=file_path,
            format=format
        )
        
        validate_file_path(str(file_path), "save")
        validate_format(format, self.SUPPORTED_FORMATS, "save")
        
        if not self._document:
            raise DocumentSaveError(
                "No document to save. Create or load a document first.",
                context
            )
        
        try:
            file_path = Path(file_path)
            
            # Ensure directory exists
            file_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Save in requested format
            if format == "docx":
                self._document.save(str(file_path))
            elif format == "pdf":
                # For PDF conversion, we need additional libraries
                self._save_as_pdf(file_path)
            elif format == "txt":
                self._save_as_text(file_path)
            else:
                # Default to docx
                self._document.save(str(file_path))
            
            return str(file_path)
            
        except Exception as e:
            raise DocumentSaveError(
                f"Failed to save document: {e}",
                context
            )
    
    def _save_as_pdf(self, file_path: Path) -> None:
        """Save document as PDF."""
        # This is a placeholder - in production, you'd use a proper PDF converter
        # like docx2pdf or unoconv
        raise NotImplementedError(
            "PDF export requires additional setup. "
            "Install docx2pdf: pip install docx2pdf"
        )
    
    def _save_as_text(self, file_path: Path) -> None:
        """Save document as plain text."""
        text_content = []
        for paragraph in self._document.paragraphs:
            text_content.append(paragraph.text)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(text_content))
    
    def _resolve_template_path(self, template: Union[str, Path]) -> Path:
        """Resolve template path."""
        template = str(template)
        
        # If it's an absolute path or relative path, use as-is
        if os.path.isabs(template) or template.startswith('.'):
            return Path(template)
        
        # Check in template directory
        if self.template_dir:
            template_path = Path(self.template_dir) / template
            if template_path.exists():
                return template_path
        
        # Check common extensions
        for ext in ['.docx', '.dotx', '.doc']:
            test_path = Path(template + ext)
            if test_path.exists():
                return test_path
        
        # Return as-is (will fail if not found)
        return Path(template)
    
    def get_document(self) -> Optional['Document']:
        """Get the current document object."""
        return self._document
    
    def clear(self) -> None:
        """Clear the current document."""
        self._document = None
    
    def is_initialized(self) -> bool:
        """Check if document is initialized."""
        return self._document is not None